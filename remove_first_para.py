from docx import Document
from docx.oxml import OxmlElement

doc = Document('final_verified_document.docx')

# Print first 5 paragraphs to understand the structure
print("=== First 5 paragraphs ===")
for i, para in enumerate(doc.paragraphs[:5]):
    print(f"Paragraph {i}: Style={para.style.name}, Text=\"{para.text}\"")
print()

# Check if first paragraph is "Integration" with Title style
if len(doc.paragraphs) > 0:
    first_para = doc.paragraphs[0]
    if first_para.text.strip() == "Integration" and first_para.style.name == "Title":
        print("Found redundant 'Integration' title at paragraph 0")
        print("Removing it...")
        
        # Remove the paragraph
        p_element = first_para._element
        p_element.getparent().remove(p_element)
        
        # Save the document
        doc.save('fixed_document.docx')
        print("Saved as 'fixed_document.docx'")
        
        # Verify the fix
        doc2 = Document('fixed_document.docx')
        print("\n=== First 5 paragraphs after fix ===")
        for i, para in enumerate(doc2.paragraphs[:5]):
            print(f"Paragraph {i}: Style={para.style.name}, Text=\"{para.text}\"")
    else:
        print(f"First paragraph is not redundant: Style={first_para.style.name}, Text=\"{first_para.text}\"")

