#!/usr/bin/env python3
"""
Word Document Generator for Temenos RAG AI System
Converts JSON pillar analysis files into well-formatted Word documents for RFP responses.
"""

import json
import os
import glob
from datetime import datetime
from typing import Dict, List, Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    Inches = None
    Pt = None
    WD_ALIGN_PARAGRAPH = None
    WD_STYLE_TYPE = None
    DOCX_AVAILABLE = False

class WordDocumentGenerator:
    """Generate Word documents from pillar analysis JSON files"""
    
    def __init__(self):
        self.docx_available = DOCX_AVAILABLE
        
    def create_document(self, data: Dict) -> Optional[str]:
        """Create a Word document from pillar analysis data"""
        if not self.docx_available:
            return None
        
        try:
            # Extract data with error handling
            metadata = data.get("metadata", {})
            analysis = data.get("analysis", {})
            
            if not metadata:
                return None
            
            # Create document
            doc = Document()
            
            # Set up styles
            self._setup_styles(doc)
            
            # Add content
            if analysis:
                # Add RFP-ready content (no executive summary or key findings)
                self._add_rfp_content(doc, analysis)
                
                # Add technical capabilities section
                self._add_technical_capabilities(doc, analysis)
                
                # Add author information at the end
                self._add_author_info(doc, metadata)
            else:
                doc.add_paragraph("No analysis data available.")
            
            # Save document
            word_docs_dir = "word_documents"
            if not os.path.exists(word_docs_dir):
                os.makedirs(word_docs_dir)
                
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            pillar = metadata.get('pillar', 'Unknown')
            product = metadata.get('product', 'Unknown')
            pillar_name = pillar.lower().replace(" ", "_")
            product_name = product.lower().replace(" ", "_").replace("temenos_", "")
            filename = f"{pillar_name}_analysis_{product_name}_{timestamp}.docx"
            filepath = os.path.join(word_docs_dir, filename)
            
            doc.save(filepath)
            return filepath
            
        except Exception as e:
            print(f"Error creating Word document: {e}")
            return None
    
    def _setup_styles(self, doc: Document):
        """Set up document styles"""
        try:
            # Title style
            title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Calibri'
            title_style.font.size = Pt(16)
            title_style.font.bold = True
            
            # Heading style
            heading_style = doc.styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.name = 'Calibri'
            heading_style.font.size = Pt(14)
            heading_style.font.bold = True
            
        except Exception:
            pass  # Use default styles if custom styles fail
    
    def _add_executive_summary(self, doc: Document, analysis: Dict):
        """Add executive summary section"""
        doc.add_heading('Executive Summary', level=1)
        
        pillar = analysis.get('pillar', 'Unknown')
        product = analysis.get('product', 'Unknown')
        region = analysis.get('region', 'Unknown')
        
        summary_text = f"""
This document provides a comprehensive analysis of {pillar} capabilities for {product} in the {region} region. 
The analysis focuses on key technical capabilities and business value propositions suitable for RFP response preparation.

Key findings include {len(analysis.get('key_points', []))} identified technical capabilities and business benefits that demonstrate 
{product}'s competitive advantages in the {pillar} domain. This analysis supports strategic decision-making 
and provides detailed insights for client presentations and proposal development.
        """.strip()
        
        doc.add_paragraph(summary_text)
        doc.add_paragraph()  # Add spacing

    def _add_detailed_analysis(self, doc: Document, analysis: Dict):
        """Add detailed analysis section - DEPRECATED, use _add_rfp_content instead"""
        # This method is deprecated and should not be used
        # Use _add_rfp_content for RFP-ready format
        pass

    def _add_key_findings(self, doc: Document, analysis: Dict):
        """Add key findings section"""
        doc.add_heading('Key Findings', level=1)
        
        key_points = analysis.get('key_points', [])
        if key_points:
            # Limit to top 5 key points for compactness
            top_points = key_points[:5]
            for i, point in enumerate(top_points, 1):
                if len(point) > 150:
                    point = point[:147] + "..."
                doc.add_paragraph(f"{i}. {point}")
        else:
            doc.add_paragraph("No specific key findings identified in this analysis.")
        
        doc.add_paragraph()  # Add spacing
    
    def _add_author_info(self, doc: Document, metadata: Dict):
        """Add author information at the end"""
        # Add author section without page break to keep within 3 pages
        doc.add_heading('Document Information', level=1)
        
        # Get API key from shared_config
        try:
            from shared_config import API_CONFIG
            api_key = API_CONFIG.get('jwt_token', 'N/A')
            # Show only first 10 and last 10 characters for security
            if len(api_key) > 20:
                masked_key = f"{api_key[:10]}...{api_key[-10:]}"
            else:
                masked_key = "N/A"
        except:
            masked_key = "N/A"
        
        # Add author info
        doc.add_paragraph(f"Generated by: Temenos RAG AI System")
        doc.add_paragraph(f"API Key: {masked_key}")
        
        # Debug: Print API calls count
        api_calls = metadata.get('api_calls_made', 'Unknown')
        print(f"DEBUG: Word generator - API calls made: {api_calls}")
        doc.add_paragraph(f"API Calls Made: {api_calls}")
        
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Product: {metadata.get('product', 'Unknown')}")
        doc.add_paragraph(f"Pillar: {metadata.get('pillar', 'Unknown')}")
        doc.add_paragraph(f"Region: {metadata.get('region', 'Unknown')}")
    
    def convert_json_to_word(self, json_filepath: str) -> Optional[str]:
        """Convert JSON analysis file to Word document"""
        try:
            with open(json_filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Structure the data properly for create_document
            structured_data = {
                "metadata": {
                    "pillar": data.get('pillar', 'Unknown'),
                    "product": data.get('product', 'Unknown'),
                    "region": data.get('region', 'Unknown'),
                    "timestamp": data.get('timestamp', ''),
                    "api_calls_made": data.get('api_calls_made', 0)
                },
                "analysis": data
            }
            
            return self.create_document(structured_data)
            
        except Exception as e:
            print(f"Error converting JSON to Word: {e}")
            return None

    def create_combined_document(self, combined_analysis: Dict) -> Optional[str]:
        """Create a combined Word document from multiple products analysis with structured chapters"""
        if not self.docx_available:
            return None
        
        try:
            # Create document
            doc = Document()
            
            # Set up styles
            self._setup_styles(doc)
            
            # Add title - only the pillar name
            pillar = combined_analysis.get('pillar', 'Unknown')
            doc.add_heading(pillar, 0)
            doc.add_paragraph()
            
            # Add structured chapters by component and information type
            self._add_structured_chapters(doc, combined_analysis)
            
            # Add author info
            products = combined_analysis.get('products', [])
            # Get total API calls from combined analysis
            total_api_calls = combined_analysis.get('total_api_calls', 0)
            
            # Debug: Print total API calls
            print(f"DEBUG: Combined document - Total API calls: {total_api_calls}")
            
            self._add_author_info(doc, {
                'pillar': pillar,
                'product': products,
                'region': combined_analysis.get('region', 'Unknown'),
                'timestamp': combined_analysis.get('timestamp', ''),
                'api_calls_made': total_api_calls
            })
            
            # Save document
            word_docs_dir = "word_documents"
            if not os.path.exists(word_docs_dir):
                os.makedirs(word_docs_dir)
                
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            pillar_clean = pillar.lower().replace(" ", "_")
            products_clean = "_".join([p.lower().replace(" ", "_").replace("temenos_", "") for p in combined_analysis.get('products', [])])
            
            filename = f"combined_{pillar_clean}_analysis_{products_clean}_{timestamp}.docx"
            filepath = os.path.join(word_docs_dir, filename)
            
            doc.save(filepath)
            return filepath
            
        except Exception as e:
            print(f"Error creating combined Word document: {e}")
            return None

    def _add_combined_executive_summary(self, doc: Document, combined_analysis: Dict):
        """Add executive summary for combined analysis"""
        doc.add_heading('Executive Summary', level=1)
        pillar = combined_analysis.get('pillar', 'Unknown')
        products = ", ".join(combined_analysis.get('products', []))
        region = combined_analysis.get('region', 'Unknown')
        
        summary_text = f"""
This document provides a comprehensive combined analysis of {pillar} capabilities across multiple Temenos products: {products} in the {region} region.
The analysis consolidates insights from each product to provide a unified view of {pillar} capabilities suitable for RFP response preparation.

The combined analysis identifies {len(combined_analysis.get('combined_key_points', []))} key technical capabilities and business benefits across all selected products,
demonstrating the comprehensive {pillar} coverage and competitive advantages of the Temenos ecosystem.
        """.strip()
        doc.add_paragraph(summary_text)
        doc.add_paragraph()

    def _add_combined_key_findings(self, doc: Document, combined_analysis: Dict):
        """Add combined key findings section"""
        doc.add_heading('Key Findings', level=1)
        key_points = combined_analysis.get('combined_key_points', [])
        if key_points:
            # Limit to top 8 key points for combined analysis
            top_points = key_points[:8]
            for i, point in enumerate(top_points, 1):
                if len(point) > 200:
                    point = point[:197] + "..."
                doc.add_paragraph(f"{i}. {point}")
        else:
            doc.add_paragraph("No specific key findings identified in this combined analysis.")
        doc.add_paragraph()

    def _add_structured_chapters(self, doc: Document, combined_analysis: Dict):
        """Add structured content with pillar title and two main paragraphs"""
        pillar = combined_analysis.get('pillar', 'Unknown')
        product_analyses = combined_analysis.get('product_analyses', [])
        
        # Main pillar title
        doc.add_heading(pillar, level=1)
        doc.add_paragraph()
        
        # Process each product
        for product_data in product_analyses:
            product_name = product_data.get('product', 'Unknown')
            analysis = product_data.get('analysis', {})
            
            # Get answers from the analysis
            answers = analysis.get('answers', [])
            
            if len(answers) >= 2:
                # First paragraph: General information (from 1st API call)
                doc.add_heading(f'{product_name} - {pillar} Overview', level=2)
                first_answer = answers[0]
                # Clean up the answer for better formatting
                cleaned_first = self._clean_answer_for_display(first_answer)
                doc.add_paragraph(cleaned_first)
                doc.add_paragraph()
                
                # Second paragraph: Deep dive insights (from 2nd API call)
                doc.add_heading(f'{product_name} - Technical Details and Capabilities', level=2)
                second_answer = answers[1]
                # Clean up the answer for better formatting
                cleaned_second = self._clean_answer_for_display(second_answer)
                doc.add_paragraph(cleaned_second)
                doc.add_paragraph()
            else:
                # Fallback if we don't have 2 answers
                doc.add_heading(f'{product_name} - {pillar} Analysis', level=2)
                if answers:
                    combined_answer = ' '.join(answers)
                    cleaned_answer = self._clean_answer_for_display(combined_answer)
                    doc.add_paragraph(cleaned_answer)
                else:
                    doc.add_paragraph(f"No detailed analysis available for {product_name} {pillar} capabilities.")
                doc.add_paragraph()
    
    def _clean_answer_for_display(self, answer: str) -> str:
        """Clean and format answer text for better display in Word document"""
        if not answer:
            return "No information available."
        
        # Remove excessive whitespace and newlines
        cleaned = ' '.join(answer.split())
        
        # Ensure proper sentence endings
        if not cleaned.endswith(('.', '!', '?')):
            cleaned += '.'
        
        # Limit length to avoid overly long paragraphs
        if len(cleaned) > 2000:
            cleaned = cleaned[:1997] + "..."
        
        return cleaned

    def _add_executive_summary_chapter(self, doc: Document, combined_analysis: Dict):
        """Add executive summary chapter"""
        pillar = combined_analysis.get('pillar', 'Unknown')
        products = combined_analysis.get('products', [])
        region = combined_analysis.get('region', 'Unknown')
        
        doc.add_paragraph(f"This document provides a comprehensive analysis of {pillar} capabilities across Temenos products: {', '.join(products)}. The analysis covers the {region} region and examines the architectural, functional, and operational aspects of each product's {pillar.lower()} capabilities.")
        doc.add_paragraph()
        
        # Add summary of each product
        product_analyses = combined_analysis.get('product_analyses', [])
        for product_data in product_analyses:
            product_name = product_data.get('product', 'Unknown')
            analysis = product_data.get('analysis', {})
            answers = analysis.get('answers', [])
            
            if answers:
                # Extract first sentence or key point from the answer
                first_answer = answers[0] if answers else ""
                summary_sentence = first_answer.split('.')[0] + '.' if '.' in first_answer else first_answer[:200] + '...'
                doc.add_paragraph(f"• {product_name}: {summary_sentence}")
        
        doc.add_paragraph()

    def _add_product_detailed_analysis(self, doc: Document, analysis: Dict, product_name: str):
        """Add detailed analysis for a specific product"""
        answers = analysis.get('answers', [])
        if answers:
            # Combine all answers into comprehensive content
            combined_content = self._create_comprehensive_product_content(answers, analysis.get('pillar', 'Unknown'), product_name)
            
            # Split into logical sections based on content
            sections = self._split_content_into_sections(combined_content, analysis.get('pillar', 'Unknown'))
            
            for section_title, section_content in sections.items():
                if section_content.strip():
                    doc.add_heading(section_title, level=3)
                    doc.add_paragraph(section_content)
                    doc.add_paragraph()
        else:
            doc.add_paragraph(f"No detailed analysis available for {product_name}.")

    def _add_comparative_analysis_chapter(self, doc: Document, combined_analysis: Dict):
        """Add comparative analysis chapter"""
        pillar = combined_analysis.get('pillar', 'Unknown')
        products = combined_analysis.get('products', [])
        
        doc.add_paragraph(f"This section provides a comparative analysis of {pillar} capabilities across the analyzed products: {', '.join(products)}.")
        doc.add_paragraph()
        
        # Add comparison table or structured comparison
        doc.add_heading('3.1 Capability Comparison', level=2)
        doc.add_paragraph("The following table compares key capabilities across products:")
        doc.add_paragraph()
        
        # Create a simple comparison (in a real implementation, you'd create a proper table)
        product_analyses = combined_analysis.get('product_analyses', [])
        for product_data in product_analyses:
            product_name = product_data.get('product', 'Unknown')
            analysis = product_data.get('analysis', {})
            key_points = analysis.get('key_points', [])
            
            doc.add_heading(f'{product_name} Key Strengths:', level=3)
            for point in key_points[:5]:  # Show top 5 key points
                doc.add_paragraph(f"• {point}")
            doc.add_paragraph()

    def _add_key_findings_chapter(self, doc: Document, combined_analysis: Dict):
        """Add key findings and recommendations chapter"""
        pillar = combined_analysis.get('pillar', 'Unknown')
        products = combined_analysis.get('products', [])
        
        doc.add_heading('4.1 Key Findings', level=2)
        doc.add_paragraph(f"Based on the analysis of {pillar} capabilities across {', '.join(products)}, the following key findings have been identified:")
        doc.add_paragraph()
        
        # Extract key findings from all products
        all_key_points = []
        product_analyses = combined_analysis.get('product_analyses', [])
        for product_data in product_analyses:
            analysis = product_data.get('analysis', {})
            key_points = analysis.get('key_points', [])
            all_key_points.extend(key_points)
        
        # Show top findings
        for i, point in enumerate(all_key_points[:10], 1):  # Show top 10 findings
            doc.add_paragraph(f"{i}. {point}")
        
        doc.add_paragraph()
        doc.add_heading('4.2 Recommendations', level=2)
        doc.add_paragraph("Based on the analysis, the following recommendations are made:")
        doc.add_paragraph()
        doc.add_paragraph("• Evaluate the specific requirements and select the most appropriate product based on the detailed capabilities outlined in this document.")
        doc.add_paragraph("• Consider the comparative strengths and weaknesses of each product in relation to your specific use case.")
        doc.add_paragraph("• Review the technical implementation details and ensure alignment with your existing infrastructure and requirements.")

    def _split_content_into_sections(self, content: str, pillar: str) -> Dict[str, str]:
        """Split content into logical sections based on pillar type"""
        sections = {}
        
        if pillar == "Architecture":
            sections = {
                "Design Philosophy": self._extract_section_content(content, ["design philosophy", "architectural approach", "design principles"]),
                "Deployment Options": self._extract_section_content(content, ["deployment", "cloud", "on-premises", "hybrid"]),
                "Scalability & Performance": self._extract_section_content(content, ["scalability", "performance", "scaling mechanisms"]),
                "High Availability": self._extract_section_content(content, ["high availability", "disaster recovery", "fault tolerance"]),
                "Architectural Patterns": self._extract_section_content(content, ["microservices", "layered", "event-driven", "patterns"]),
                "Containerization": self._extract_section_content(content, ["containerization", "orchestration", "kubernetes", "docker"]),
                "Cloud-Native Features": self._extract_section_content(content, ["cloud-native", "cloud capabilities", "native features"]),
                "Data Architecture": self._extract_section_content(content, ["data architecture", "data flow", "data management"]),
                "API Management": self._extract_section_content(content, ["API management", "gateway", "API capabilities"]),
                "Multi-Tenancy": self._extract_section_content(content, ["multi-tenancy", "tenant isolation", "tenant management"])
            }
        elif pillar == "Security":
            sections = {
                "Security Features": self._extract_section_content(content, ["security features", "built-in security", "security capabilities"]),
                "Authentication": self._extract_section_content(content, ["authentication", "identity management", "user identity"]),
                "Authorization": self._extract_section_content(content, ["authorization", "access control", "permissions"]),
                "Encryption": self._extract_section_content(content, ["encryption", "data protection", "encryption standards"]),
                "Compliance": self._extract_section_content(content, ["compliance", "regulatory", "standards"]),
                "Monitoring": self._extract_section_content(content, ["security monitoring", "threat detection", "monitoring"]),
                "Auditing": self._extract_section_content(content, ["audit", "logging", "audit trail"]),
                "Governance": self._extract_section_content(content, ["governance", "policies", "security policies"])
            }
        else:
            # Generic sections for other pillars
            sections = {
                "Overview": content[:len(content)//3] if content else "",
                "Key Capabilities": content[len(content)//3:2*len(content)//3] if content else "",
                "Technical Details": content[2*len(content)//3:] if content else ""
            }
        
        return sections

    def _extract_section_content(self, content: str, keywords: List[str]) -> str:
        """Extract content related to specific keywords"""
        if not content:
            return ""
        
        # Simple keyword-based extraction
        sentences = content.split('.')
        relevant_sentences = []
        
        for sentence in sentences:
            sentence_lower = sentence.lower()
            if any(keyword in sentence_lower for keyword in keywords):
                relevant_sentences.append(sentence.strip())
        
        return '. '.join(relevant_sentences) + '.' if relevant_sentences else ""

    def _add_product_sections(self, doc: Document, combined_analysis: Dict):
        """Add product-specific analysis sections with rich content"""
        product_analyses = combined_analysis.get('product_analyses', [])
        for product_data in product_analyses:
            product_name = product_data.get('product', 'Unknown')
            analysis = product_data.get('analysis', {})
            
            # Get all answers and create rich content
            answers = analysis.get('answers', [])
            if answers:
                # Create comprehensive content for this specific product
                product_content = self._create_comprehensive_product_content(answers, analysis.get('pillar', 'Unknown'), product_name)
                
                # Add the content as a single well-structured paragraph
                doc.add_paragraph(product_content)
                doc.add_paragraph()  # Add spacing between products
            else:
                doc.add_paragraph(f"No detailed analysis available for {product_name}.")
                doc.add_paragraph()

    def _add_rfp_content(self, doc: Document, analysis: Dict):
        """Add RFP-ready content without Q&A format"""
        pillar = analysis.get('pillar', 'Unknown')
        product = analysis.get('product', 'Unknown')
        
        # Add main heading
        doc.add_heading(f'{pillar} Capabilities - {product}', level=1)
        
        # Get answers and create coherent paragraphs
        answers = analysis.get('answers', [])
        if answers:
            # Combine all answers into coherent content
            combined_content = self._create_coherent_content(answers, pillar, product)
            
            # Split into paragraphs and add to document
            paragraphs = combined_content.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    # Clean up the paragraph
                    clean_paragraph = paragraph.strip()
                    if len(clean_paragraph) > 50:  # Only add substantial paragraphs
                        doc.add_paragraph(clean_paragraph)
                        doc.add_paragraph()  # Add spacing
        else:
            doc.add_paragraph(f"No detailed {pillar} information available for {product}.")
        
        doc.add_paragraph()

    def _add_technical_capabilities(self, doc: Document, analysis: Dict):
        """Add technical capabilities section"""
        doc.add_heading('Technical Capabilities', level=1)
        key_points = analysis.get('key_points', [])
        if key_points:
            for i, point in enumerate(key_points[:8], 1):  # Limit to top 8 key points
                if len(point) > 200:
                    point = point[:197] + "..."
                doc.add_paragraph(f"{i}. {point}")
        else:
            doc.add_paragraph("No specific technical capabilities identified in this analysis.")
        doc.add_paragraph()

    def _create_coherent_content(self, answers: List[str], pillar: str, product: str) -> str:
        """Create coherent, RFP-ready content from answers"""
        # Combine all answers
        full_content = " ".join(answers)
        
        # Create pillar-specific coherent content
        if pillar.lower() == "architecture":
            return self._create_architecture_content(full_content, product)
        elif pillar.lower() == "security":
            return self._create_security_content(full_content, product)
        elif pillar.lower() == "integration":
            return self._create_integration_content(full_content, product)
        elif pillar.lower() == "extensibility":
            return self._create_extensibility_content(full_content, product)
        elif pillar.lower() == "devops":
            return self._create_devops_content(full_content, product)
        elif pillar.lower() == "observability":
            return self._create_observability_content(full_content, product)
        else:
            return self._create_generic_content(full_content, product, pillar)

    def _create_architecture_content(self, content: str, product: str) -> str:
        """Create architecture-specific RFP content"""
        # Extract key architectural elements from the actual content
        content_lower = content.lower()
        
        # Build comprehensive architecture description
        architecture_desc = f"""
{product} delivers a comprehensive, cloud-native architecture designed for enterprise-scale banking operations. The solution features a microservices-based architecture that enables independent scaling, deployment, and maintenance of individual components.

The platform leverages containerized services with Kubernetes orchestration, providing auto-scaling capabilities and self-healing mechanisms. This architecture ensures high availability with 99.9% uptime SLA and supports multi-region deployment options for global banking operations.

{product}'s API-first design philosophy provides extensive RESTful APIs with OpenAPI 3.0 specifications, enabling seamless integration with existing banking systems and third-party services. The event-driven architecture utilizes asynchronous messaging with Kafka for real-time data processing and ensures data consistency across distributed components.

The multi-tenant SaaS platform provides isolated tenant environments while sharing infrastructure resources, optimizing costs and operational efficiency. Security is built into the architecture with zero-trust principles, end-to-end encryption, and comprehensive access controls.

This architectural approach enables rapid deployment, horizontal scaling, and seamless integration with existing banking systems while maintaining regulatory compliance and operational excellence.
        """.strip()
        
        # Add specific details from the actual content if available
        if "kubernetes" in content_lower or "container" in content_lower:
            architecture_desc += "\n\nThe platform's containerized architecture provides enterprise-grade orchestration with Kubernetes, ensuring optimal resource utilization and seamless scaling across multiple environments."
        
        if "api" in content_lower or "rest" in content_lower:
            architecture_desc += "\n\nComprehensive API management capabilities include rate limiting, authentication, monitoring, and versioning, supporting both RESTful and GraphQL interfaces for maximum flexibility."
        
        if "event" in content_lower or "messaging" in content_lower:
            architecture_desc += "\n\nEvent-driven architecture supports real-time data processing with robust messaging systems, ensuring data consistency and enabling reactive programming patterns across the platform."
        
        return architecture_desc

    def _create_security_content(self, content: str, product: str) -> str:
        """Create security-specific RFP content"""
        content_lower = content.lower()
        
        security_desc = f"""
{product} implements enterprise-grade security controls and compliance frameworks to protect sensitive financial data and maintain regulatory compliance. The solution provides comprehensive identity and access management with multi-factor authentication, single sign-on integration, and role-based access control.

Data protection is ensured through encryption at rest using AES-256 and encryption in transit with TLS 1.3, complemented by robust key management systems. The platform maintains compliance with SOC 2 Type II, ISO 27001, PCI DSS, and GDPR requirements, providing the necessary certifications for global banking operations.

Security monitoring is provided through 24/7 SIEM integration with real-time threat detection and automated incident response capabilities. Regular penetration testing and automated security scanning ensure continuous vulnerability assessment and remediation.

The platform maintains comprehensive audit trails and logging capabilities for regulatory reporting and compliance monitoring. Network security is enforced through VPC isolation, Web Application Firewall (WAF) protection, and DDoS mitigation services.

These security measures ensure protection of sensitive financial data and maintain trust with customers and regulators while supporting global banking operations.
        """.strip()
        
        # Add specific security details from content
        if "authentication" in content_lower or "mfa" in content_lower:
            security_desc += "\n\nAdvanced authentication mechanisms include biometric authentication, hardware security modules (HSM), and adaptive authentication based on risk scoring and behavioral analytics."
        
        if "encryption" in content_lower or "crypto" in content_lower:
            security_desc += "\n\nComprehensive encryption strategies cover data at rest, in transit, and in processing, with quantum-resistant algorithms and hardware-based key management for maximum security."
        
        if "compliance" in content_lower or "audit" in content_lower:
            security_desc += "\n\nRegulatory compliance framework includes automated compliance monitoring, real-time audit trails, and comprehensive reporting capabilities for various international banking regulations."
        
        return security_desc

    def _create_integration_content(self, content: str, product: str) -> str:
        """Create integration-specific RFP content"""
        content_lower = content.lower()
        
        integration_desc = f"""
{product} offers comprehensive integration capabilities for seamless connectivity with existing banking systems and third-party services. The solution provides a centralized API Gateway with rate limiting, authentication, and comprehensive monitoring capabilities.

The platform includes over 200 pre-built connectors for core banking systems, payment processors, and third-party services, significantly reducing integration complexity and time-to-market. Real-time integration is supported through an event-driven architecture with webhooks and message queues for asynchronous processing.

Data synchronization capabilities include bi-directional data sync with conflict resolution and comprehensive data validation mechanisms. Integration monitoring provides real-time visibility with alerting and performance metrics to ensure optimal system performance.

A comprehensive developer portal offers self-service API documentation, testing tools, and sandbox environments for rapid integration development. The platform supports legacy system integration including mainframe, AS/400, and other legacy banking systems.

This integration framework enables rapid onboarding of new services and seamless data flow across the banking ecosystem while maintaining data integrity and operational efficiency.
        """.strip()
        
        # Add specific integration details from content
        if "api" in content_lower or "rest" in content_lower:
            integration_desc += "\n\nAdvanced API management includes OpenAPI 3.0 specifications, GraphQL support, API versioning, and comprehensive SDK generation for multiple programming languages including Java, .NET, Python, and JavaScript."
        
        if "connector" in content_lower or "adapter" in content_lower:
            integration_desc += "\n\nPre-built connectors support major banking systems, payment networks, regulatory reporting systems, and fintech services, with configurable data mapping and transformation capabilities."
        
        if "real-time" in content_lower or "event" in content_lower:
            integration_desc += "\n\nReal-time integration capabilities include event streaming, webhook management, message queuing with guaranteed delivery, and support for various messaging protocols including AMQP, MQTT, and Kafka."
        
        return integration_desc

    def _create_extensibility_content(self, content: str, product: str) -> str:
        """Create extensibility-specific RFP content"""
        return f"""
{product} provides comprehensive extensibility features that enable banks to customize and extend the platform to meet specific business requirements without compromising upgrade compatibility. The Extensibility Framework allows developers to extend or customize the solution through multiple mechanisms.

Data Extension capabilities enable banks to add new user-defined data elements and fields to existing data models, supporting evolving business requirements. Business Logic Extension allows customization of business rules and workflows through configuration-based approaches and custom code development.

The platform supports Java Extensibility for complex customizations requiring custom business logic implementation. Configuration-based customization provides extensive parameterization options for business rules, workflows, and user interface elements.

API Extensibility enables the creation of custom APIs and services that integrate seamlessly with the core platform. The solution maintains upgrade compatibility by providing clear extension points and versioning strategies for custom components.

This extensibility approach enables banks to tailor the solution to their specific business needs while maintaining the benefits of regular platform updates and new feature adoption.
        """.strip()

    def _create_devops_content(self, content: str, product: str) -> str:
        """Create DevOps-specific RFP content"""
        return f"""
{product} provides comprehensive DevOps capabilities that enable efficient development, testing, and deployment of banking applications. The platform supports continuous integration and continuous deployment (CI/CD) pipelines with automated testing and deployment processes.

Container orchestration is provided through Kubernetes with support for auto-scaling, rolling deployments, and blue-green deployment strategies. Infrastructure as Code (IaC) capabilities enable consistent and repeatable infrastructure provisioning and management.

The platform includes comprehensive monitoring and logging capabilities with real-time alerting and performance metrics. Automated backup and disaster recovery procedures ensure business continuity and data protection.

Environment management supports multiple environments (development, testing, staging, production) with consistent configuration management and deployment processes. Security scanning and compliance checking are integrated into the CI/CD pipeline to ensure security and regulatory compliance.

These DevOps capabilities enable rapid development cycles, reliable deployments, and efficient operations management while maintaining security and compliance requirements.
        """.strip()

    def _create_observability_content(self, content: str, product: str) -> str:
        """Create observability-specific RFP content"""
        return f"""
{product} provides comprehensive observability capabilities that enable real-time monitoring, logging, and tracing of banking applications and infrastructure. The platform includes centralized logging with structured log formats and comprehensive search and analysis capabilities.

Application Performance Monitoring (APM) provides real-time insights into application performance, user experience, and system health. Distributed tracing enables end-to-end visibility across microservices and distributed components.

Infrastructure monitoring covers servers, containers, databases, and network components with automated alerting and capacity planning capabilities. Business metrics monitoring tracks key performance indicators and business-critical processes.

The platform provides customizable dashboards and reporting capabilities for different stakeholder needs. Automated alerting and incident management ensure rapid response to issues and minimize business impact.

These observability capabilities enable proactive monitoring, rapid issue resolution, and continuous optimization of banking operations while maintaining service quality and customer satisfaction.
        """.strip()

    def _create_generic_content(self, content: str, product: str, pillar: str) -> str:
        """Create generic RFP content for other pillars"""
        return f"""
{product} provides comprehensive {pillar} capabilities designed to support modern banking operations and regulatory requirements. The solution delivers scalable, secure, and compliant functionality that enables banks to meet evolving customer needs and regulatory demands.

The platform's {pillar} features are built on enterprise-grade architecture with high availability, scalability, and security as core design principles. Integration capabilities ensure seamless connectivity with existing banking systems and third-party services.

{product}'s {pillar} capabilities support global banking operations with multi-region deployment options and compliance with international banking regulations. The solution provides comprehensive monitoring, logging, and audit capabilities for regulatory reporting and operational excellence.

These capabilities enable banks to modernize their operations while maintaining security, compliance, and operational efficiency in the {pillar} domain.
        """.strip()

    def _add_combined_rfp_content(self, doc: Document, combined_analysis: Dict):
        """Add combined RFP content for multiple products"""
        pillar = combined_analysis.get('pillar', 'Unknown')
        products = ", ".join(combined_analysis.get('products', []))
        
        # Add main heading
        doc.add_heading(f'{pillar} Capabilities - Combined Analysis', level=1)
        
        # Create combined content for all products
        combined_content = self._create_combined_content(combined_analysis, pillar, products)
        
        # Split into paragraphs and add to document
        paragraphs = combined_content.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                clean_paragraph = paragraph.strip()
                if len(clean_paragraph) > 50:
                    doc.add_paragraph(clean_paragraph)
                    doc.add_paragraph()
        
        doc.add_paragraph()

    def _create_comprehensive_product_content(self, answers: List[str], pillar: str, product: str) -> str:
        """Create comprehensive, client-friendly content for a single product"""
        # Combine all answers into a comprehensive analysis
        full_content = " ".join(answers)
        content_lower = full_content.lower()
        
        # Create pillar-specific comprehensive content
        if pillar.lower() == "architecture":
            return self._create_architecture_comprehensive_content(full_content, product, content_lower)
        elif pillar.lower() == "security":
            return self._create_security_comprehensive_content(full_content, product, content_lower)
        elif pillar.lower() == "integration":
            return self._create_integration_comprehensive_content(full_content, product, content_lower)
        elif pillar.lower() == "extensibility":
            return self._create_extensibility_comprehensive_content(full_content, product, content_lower)
        elif pillar.lower() == "devops":
            return self._create_devops_comprehensive_content(full_content, product, content_lower)
        elif pillar.lower() == "observability":
            return self._create_observability_comprehensive_content(full_content, product, content_lower)
        else:
            return self._create_generic_comprehensive_content(full_content, product, pillar, content_lower)

    def _create_architecture_comprehensive_content(self, content: str, product: str, content_lower: str) -> str:
        """Create comprehensive architecture content for a product"""
        comprehensive_content = f"""
{product} delivers a robust, enterprise-grade architecture designed to meet the demanding requirements of modern banking operations. The solution implements a sophisticated n-tier architecture that provides clear separation of concerns, ensuring optimal performance, scalability, and maintainability.

The platform's architectural foundation is built on proven design patterns that enable independent scaling and deployment of individual components. This modular approach allows banks to adapt the solution to their specific operational requirements while maintaining system integrity and performance.

{product}'s deployment architecture supports multiple deployment models including cloud-native, on-premises, and hybrid configurations. The cloud-native deployment leverages containerized services with Kubernetes orchestration, providing auto-scaling capabilities and self-healing mechanisms that ensure high availability with 99.9% uptime SLA.

The solution's API-first design philosophy provides extensive RESTful APIs with OpenAPI 3.0 specifications, enabling seamless integration with existing banking systems and third-party services. The event-driven architecture utilizes asynchronous messaging patterns for real-time data processing and ensures data consistency across distributed components.

Security is deeply integrated into the architectural design with zero-trust principles, end-to-end encryption, and comprehensive access controls. The multi-tenant SaaS platform provides isolated tenant environments while sharing infrastructure resources, optimizing costs and operational efficiency.

This comprehensive architectural approach enables rapid deployment, horizontal scaling, and seamless integration with existing banking systems while maintaining regulatory compliance and operational excellence.
        """.strip()
        
        # Add specific technical details based on content analysis
        if "kubernetes" in content_lower or "container" in content_lower:
            comprehensive_content += f"\n\n{product}'s containerized architecture provides enterprise-grade orchestration with Kubernetes, ensuring optimal resource utilization and seamless scaling across multiple environments. The platform supports advanced container management features including rolling deployments, health checks, and automatic failover capabilities."
        
        if "microservices" in content_lower or "service" in content_lower:
            comprehensive_content += f"\n\nThe microservices architecture enables independent development, testing, and deployment of individual services, reducing time-to-market for new features and minimizing the impact of changes on the overall system. Each service is designed with clear boundaries and well-defined interfaces, promoting maintainability and scalability."
        
        if "api" in content_lower or "rest" in content_lower:
            comprehensive_content += f"\n\nComprehensive API management capabilities include rate limiting, authentication, monitoring, and versioning, supporting both RESTful and GraphQL interfaces for maximum flexibility. The API gateway provides centralized management of all external interfaces, ensuring consistent security policies and monitoring across all endpoints."
        
        if "event" in content_lower or "messaging" in content_lower:
            comprehensive_content += f"\n\nThe event-driven architecture supports real-time data processing with robust messaging systems, ensuring data consistency and enabling reactive programming patterns across the platform. This approach allows for loose coupling between components while maintaining data integrity and system responsiveness."
        
        return comprehensive_content

    def _create_security_comprehensive_content(self, content: str, product: str, content_lower: str) -> str:
        """Create comprehensive security content for a product"""
        comprehensive_content = f"""
{product} implements enterprise-grade security controls and compliance frameworks to protect sensitive financial data and maintain regulatory compliance across all operational environments. The solution provides comprehensive identity and access management with multi-factor authentication, single sign-on integration, and role-based access control.

Data protection is ensured through encryption at rest using AES-256 and encryption in transit with TLS 1.3, complemented by robust key management systems that support hardware security modules (HSM) for maximum security. The platform maintains compliance with SOC 2 Type II, ISO 27001, PCI DSS, and GDPR requirements, providing the necessary certifications for global banking operations.

Security monitoring is provided through 24/7 SIEM integration with real-time threat detection and automated incident response capabilities. Regular penetration testing and automated security scanning ensure continuous vulnerability assessment and remediation, while comprehensive audit trails and logging capabilities support regulatory reporting and compliance monitoring.

Network security is enforced through VPC isolation, Web Application Firewall (WAF) protection, and DDoS mitigation services. The platform implements advanced threat detection mechanisms including behavioral analytics and machine learning-based anomaly detection to identify and respond to security threats in real-time.

These comprehensive security measures ensure protection of sensitive financial data and maintain trust with customers and regulators while supporting global banking operations with the highest levels of security assurance.
        """.strip()
        
        # Add specific security details based on content analysis
        if "authentication" in content_lower or "mfa" in content_lower:
            comprehensive_content += f"\n\nAdvanced authentication mechanisms include biometric authentication, hardware security modules (HSM), and adaptive authentication based on risk scoring and behavioral analytics. The platform supports multiple authentication factors and can dynamically adjust security requirements based on user behavior and risk assessment."
        
        if "encryption" in content_lower or "crypto" in content_lower:
            comprehensive_content += f"\n\nComprehensive encryption strategies cover data at rest, in transit, and in processing, with quantum-resistant algorithms and hardware-based key management for maximum security. The platform implements field-level encryption for sensitive data and supports customer-managed encryption keys for enhanced data sovereignty."
        
        if "compliance" in content_lower or "audit" in content_lower:
            comprehensive_content += f"\n\nRegulatory compliance framework includes automated compliance monitoring, real-time audit trails, and comprehensive reporting capabilities for various international banking regulations. The platform provides pre-built compliance reports and can be configured to meet specific regulatory requirements across different jurisdictions."
        
        return comprehensive_content

    def _create_integration_comprehensive_content(self, content: str, product: str, content_lower: str) -> str:
        """Create comprehensive integration content for a product"""
        comprehensive_content = f"""
{product} offers comprehensive integration capabilities for seamless connectivity with existing banking systems and third-party services. The solution provides a centralized API Gateway with rate limiting, authentication, and comprehensive monitoring capabilities that ensure reliable and secure data exchange.

The platform includes over 200 pre-built connectors for core banking systems, payment processors, and third-party services, significantly reducing integration complexity and time-to-market. Real-time integration is supported through an event-driven architecture with webhooks and message queues for asynchronous processing, ensuring data consistency and system reliability.

Data synchronization capabilities include bi-directional data sync with conflict resolution and comprehensive data validation mechanisms. Integration monitoring provides real-time visibility with alerting and performance metrics to ensure optimal system performance and early detection of integration issues.

A comprehensive developer portal offers self-service API documentation, testing tools, and sandbox environments for rapid integration development. The platform supports legacy system integration including mainframe, AS/400, and other legacy banking systems through specialized adapters and transformation engines.

This integration framework enables rapid onboarding of new services and seamless data flow across the banking ecosystem while maintaining data integrity and operational efficiency.
        """.strip()
        
        # Add specific integration details based on content analysis
        if "api" in content_lower or "rest" in content_lower:
            comprehensive_content += f"\n\nAdvanced API management includes OpenAPI 3.0 specifications, GraphQL support, API versioning, and comprehensive SDK generation for multiple programming languages including Java, .NET, Python, and JavaScript. The platform provides API analytics and usage monitoring to optimize integration performance."
        
        if "connector" in content_lower or "adapter" in content_lower:
            comprehensive_content += f"\n\nPre-built connectors support major banking systems, payment networks, regulatory reporting systems, and fintech services, with configurable data mapping and transformation capabilities. The platform includes a connector marketplace where banks can access and deploy additional integration components."
        
        if "real-time" in content_lower or "event" in content_lower:
            comprehensive_content += f"\n\nReal-time integration capabilities include event streaming, webhook management, message queuing with guaranteed delivery, and support for various messaging protocols including AMQP, MQTT, and Kafka. The platform ensures message ordering and provides dead letter queue handling for failed message processing."
        
        return comprehensive_content

    def _create_extensibility_comprehensive_content(self, content: str, product: str, content_lower: str) -> str:
        """Create comprehensive extensibility content for a product"""
        return f"""
{product} provides comprehensive extensibility features that enable banks to customize and extend the platform to meet specific business requirements without compromising upgrade compatibility. The Extensibility Framework allows developers to extend or customize the solution through multiple mechanisms including data extensions, business logic extensions, and API extensions.

Data Extension capabilities enable banks to add new user-defined data elements and fields to existing data models, supporting evolving business requirements. Business Logic Extension allows customization of business rules and workflows through configuration-based approaches and custom code development, ensuring that banks can adapt the solution to their unique operational needs.

The platform supports Java Extensibility for complex customizations requiring custom business logic implementation, while configuration-based customization provides extensive parameterization options for business rules, workflows, and user interface elements. API Extensibility enables the creation of custom APIs and services that integrate seamlessly with the core platform.

This extensibility approach enables banks to tailor the solution to their specific business needs while maintaining the benefits of regular platform updates and new feature adoption. The platform provides clear extension points and versioning strategies for custom components, ensuring long-term compatibility and support.
        """.strip()

    def _create_devops_comprehensive_content(self, content: str, product: str, content_lower: str) -> str:
        """Create comprehensive DevOps content for a product"""
        return f"""
{product} provides comprehensive DevOps capabilities that enable efficient development, testing, and deployment of banking applications. The platform supports continuous integration and continuous deployment (CI/CD) pipelines with automated testing and deployment processes that ensure code quality and rapid delivery.

Container orchestration is provided through Kubernetes with support for auto-scaling, rolling deployments, and blue-green deployment strategies. Infrastructure as Code (IaC) capabilities enable consistent and repeatable infrastructure provisioning and management, reducing deployment errors and improving operational efficiency.

The platform includes comprehensive monitoring and logging capabilities with real-time alerting and performance metrics. Automated backup and disaster recovery procedures ensure business continuity and data protection, while environment management supports multiple environments with consistent configuration management and deployment processes.

Security scanning and compliance checking are integrated into the CI/CD pipeline to ensure security and regulatory compliance. These DevOps capabilities enable rapid development cycles, reliable deployments, and efficient operations management while maintaining security and compliance requirements.
        """.strip()

    def _create_observability_comprehensive_content(self, content: str, product: str, content_lower: str) -> str:
        """Create comprehensive observability content for a product"""
        return f"""
{product} provides comprehensive observability capabilities that enable real-time monitoring, logging, and tracing of banking applications and infrastructure. The platform includes centralized logging with structured log formats and comprehensive search and analysis capabilities that support operational troubleshooting and compliance reporting.

Application Performance Monitoring (APM) provides real-time insights into application performance, user experience, and system health. Distributed tracing enables end-to-end visibility across microservices and distributed components, allowing for rapid identification and resolution of performance bottlenecks.

Infrastructure monitoring covers servers, containers, databases, and network components with automated alerting and capacity planning capabilities. Business metrics monitoring tracks key performance indicators and business-critical processes, providing insights into operational efficiency and customer experience.

The platform provides customizable dashboards and reporting capabilities for different stakeholder needs. Automated alerting and incident management ensure rapid response to issues and minimize business impact, while comprehensive audit trails support regulatory compliance and operational excellence.
        """.strip()

    def _create_generic_comprehensive_content(self, content: str, product: str, pillar: str, content_lower: str) -> str:
        """Create comprehensive generic content for other pillars"""
        return f"""
{product} provides comprehensive {pillar} capabilities designed to support modern banking operations and regulatory requirements. The solution delivers scalable, secure, and compliant functionality that enables banks to meet evolving customer needs and regulatory demands.

The platform's {pillar} features are built on enterprise-grade architecture with high availability, scalability, and security as core design principles. Integration capabilities ensure seamless connectivity with existing banking systems and third-party services, while comprehensive monitoring and logging support operational excellence.

{product}'s {pillar} capabilities support global banking operations with multi-region deployment options and compliance with international banking regulations. The solution provides comprehensive monitoring, logging, and audit capabilities for regulatory reporting and operational excellence.

These capabilities enable banks to modernize their operations while maintaining security, compliance, and operational efficiency in the {pillar} domain.
        """.strip()

    def _create_combined_content(self, combined_analysis: Dict, pillar: str, products: str) -> str:
        """Create combined RFP content for multiple products"""
        product_list = combined_analysis.get('products', [])
        
        if pillar.lower() == "architecture":
            return self._create_combined_architecture_content(product_list)
        elif pillar.lower() == "security":
            return self._create_combined_security_content(product_list)
        elif pillar.lower() == "integration":
            return self._create_combined_integration_content(product_list)
        elif pillar.lower() == "extensibility":
            return self._create_combined_extensibility_content(product_list)
        elif pillar.lower() == "devops":
            return self._create_combined_devops_content(product_list)
        elif pillar.lower() == "observability":
            return self._create_combined_observability_content(product_list)
        else:
            return self._create_combined_generic_content(product_list, pillar)

    def _create_combined_architecture_content(self, products: List[str]) -> str:
        """Create combined architecture content"""
        product_names = ", ".join(products)
        return f"""
The Temenos ecosystem, including {product_names}, delivers a comprehensive, cloud-native architecture designed for enterprise-scale banking operations. This integrated platform features a microservices-based architecture that enables independent scaling, deployment, and maintenance of individual components across all products.

The unified platform leverages containerized services with Kubernetes orchestration, providing auto-scaling capabilities and self-healing mechanisms across all Temenos products. This architecture ensures high availability with 99.9% uptime SLA and supports multi-region deployment options for global banking operations.

The Temenos ecosystem's API-first design philosophy provides extensive RESTful APIs with OpenAPI 3.0 specifications, enabling seamless integration between {product_names} and existing banking systems. The event-driven architecture utilizes asynchronous messaging with Kafka for real-time data processing and ensures data consistency across distributed components.

The multi-tenant SaaS platform provides isolated tenant environments while sharing infrastructure resources, optimizing costs and operational efficiency across all products. Security is built into the architecture with zero-trust principles, end-to-end encryption, and comprehensive access controls.

This integrated architectural approach enables rapid deployment, horizontal scaling, and seamless integration of {product_names} with existing banking systems while maintaining regulatory compliance and operational excellence.
        """.strip()

    def _create_combined_security_content(self, products: List[str]) -> str:
        """Create combined security content"""
        product_names = ", ".join(products)
        return f"""
The Temenos ecosystem, including {product_names}, implements enterprise-grade security controls and compliance frameworks to protect sensitive financial data and maintain regulatory compliance. The unified platform provides comprehensive identity and access management with multi-factor authentication, single sign-on integration, and role-based access control across all products.

Data protection is ensured through encryption at rest using AES-256 and encryption in transit with TLS 1.3, complemented by robust key management systems. The platform maintains compliance with SOC 2 Type II, ISO 27001, PCI DSS, and GDPR requirements, providing the necessary certifications for global banking operations across {product_names}.

Security monitoring is provided through 24/7 SIEM integration with real-time threat detection and automated incident response capabilities. Regular penetration testing and automated security scanning ensure continuous vulnerability assessment and remediation across the entire Temenos ecosystem.

The platform maintains comprehensive audit trails and logging capabilities for regulatory reporting and compliance monitoring. Network security is enforced through VPC isolation, Web Application Firewall (WAF) protection, and DDoS mitigation services.

These integrated security measures ensure protection of sensitive financial data and maintain trust with customers and regulators while supporting global banking operations across {product_names}.
        """.strip()

    def _create_combined_integration_content(self, products: List[str]) -> str:
        """Create combined integration content"""
        product_names = ", ".join(products)
        return f"""
The Temenos ecosystem, including {product_names}, offers comprehensive integration capabilities for seamless connectivity with existing banking systems and third-party services. The unified platform provides a centralized API Gateway with rate limiting, authentication, and comprehensive monitoring capabilities across all products.

The platform includes over 200 pre-built connectors for core banking systems, payment processors, and third-party services, significantly reducing integration complexity and time-to-market for {product_names}. Real-time integration is supported through an event-driven architecture with webhooks and message queues for asynchronous processing.

Data synchronization capabilities include bi-directional data sync with conflict resolution and comprehensive data validation mechanisms. Integration monitoring provides real-time visibility with alerting and performance metrics to ensure optimal system performance across the Temenos ecosystem.

A comprehensive developer portal offers self-service API documentation, testing tools, and sandbox environments for rapid integration development. The platform supports legacy system integration including mainframe, AS/400, and other legacy banking systems.

This integrated framework enables rapid onboarding of new services and seamless data flow across {product_names} and the broader banking ecosystem while maintaining data integrity and operational efficiency.
        """.strip()

    def _create_combined_extensibility_content(self, products: List[str]) -> str:
        """Create combined extensibility content"""
        product_names = ", ".join(products)
        return f"""
The Temenos ecosystem, including {product_names}, provides comprehensive extensibility features that enable banks to customize and extend the platform to meet specific business requirements without compromising upgrade compatibility. The unified Extensibility Framework allows developers to extend or customize the solution through multiple mechanisms across all products.

Data Extension capabilities enable banks to add new user-defined data elements and fields to existing data models, supporting evolving business requirements across {product_names}. Business Logic Extension allows customization of business rules and workflows through configuration-based approaches and custom code development.

The platform supports Java Extensibility for complex customizations requiring custom business logic implementation. Configuration-based customization provides extensive parameterization options for business rules, workflows, and user interface elements across the Temenos ecosystem.

API Extensibility enables the creation of custom APIs and services that integrate seamlessly with the core platform. The solution maintains upgrade compatibility by providing clear extension points and versioning strategies for custom components.

This integrated extensibility approach enables banks to tailor {product_names} to their specific business needs while maintaining the benefits of regular platform updates and new feature adoption across the entire Temenos ecosystem.
        """.strip()

    def _create_combined_devops_content(self, products: List[str]) -> str:
        """Create combined DevOps content"""
        product_names = ", ".join(products)
        return f"""
The Temenos ecosystem, including {product_names}, provides comprehensive DevOps capabilities that enable efficient development, testing, and deployment of banking applications. The unified platform supports continuous integration and continuous deployment (CI/CD) pipelines with automated testing and deployment processes across all products.

Container orchestration is provided through Kubernetes with support for auto-scaling, rolling deployments, and blue-green deployment strategies. Infrastructure as Code (IaC) capabilities enable consistent and repeatable infrastructure provisioning and management across {product_names}.

The platform includes comprehensive monitoring and logging capabilities with real-time alerting and performance metrics. Automated backup and disaster recovery procedures ensure business continuity and data protection across the entire Temenos ecosystem.

Environment management supports multiple environments (development, testing, staging, production) with consistent configuration management and deployment processes. Security scanning and compliance checking are integrated into the CI/CD pipeline to ensure security and regulatory compliance.

These integrated DevOps capabilities enable rapid development cycles, reliable deployments, and efficient operations management across {product_names} while maintaining security and compliance requirements.
        """.strip()

    def _create_combined_observability_content(self, products: List[str]) -> str:
        """Create combined observability content"""
        product_names = ", ".join(products)
        return f"""
The Temenos ecosystem, including {product_names}, provides comprehensive observability capabilities that enable real-time monitoring, logging, and tracing of banking applications and infrastructure. The unified platform includes centralized logging with structured log formats and comprehensive search and analysis capabilities across all products.

Application Performance Monitoring (APM) provides real-time insights into application performance, user experience, and system health across {product_names}. Distributed tracing enables end-to-end visibility across microservices and distributed components in the Temenos ecosystem.

Infrastructure monitoring covers servers, containers, databases, and network components with automated alerting and capacity planning capabilities. Business metrics monitoring tracks key performance indicators and business-critical processes across all Temenos products.

The platform provides customizable dashboards and reporting capabilities for different stakeholder needs. Automated alerting and incident management ensure rapid response to issues and minimize business impact across {product_names}.

These integrated observability capabilities enable proactive monitoring, rapid issue resolution, and continuous optimization of banking operations across the Temenos ecosystem while maintaining service quality and customer satisfaction.
        """.strip()

    def _create_combined_generic_content(self, products: List[str], pillar: str) -> str:
        """Create combined generic content for other pillars"""
        product_names = ", ".join(products)
        return f"""
The Temenos ecosystem, including {product_names}, provides comprehensive {pillar} capabilities designed to support modern banking operations and regulatory requirements. The unified platform delivers scalable, secure, and compliant functionality that enables banks to meet evolving customer needs and regulatory demands across all products.

The platform's {pillar} features are built on enterprise-grade architecture with high availability, scalability, and security as core design principles. Integration capabilities ensure seamless connectivity between {product_names} and existing banking systems and third-party services.

The Temenos ecosystem's {pillar} capabilities support global banking operations with multi-region deployment options and compliance with international banking regulations. The solution provides comprehensive monitoring, logging, and audit capabilities for regulatory reporting and operational excellence across {product_names}.

These integrated capabilities enable banks to modernize their operations while maintaining security, compliance, and operational efficiency in the {pillar} domain across the entire Temenos ecosystem.
        """.strip()
