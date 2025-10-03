#!/usr/bin/env python3
"""
Word Document Generator for Temenos RAG AI System
Converts JSON pillar analysis files into well-formatted Word documents for RFP responses.
"""

import json
import os
import glob
from datetime import datetime
from typing import Dict, List, Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    Inches = None
    Pt = None
    WD_ALIGN_PARAGRAPH = None
    WD_STYLE_TYPE = None
    DOCX_AVAILABLE = False

class WordDocumentGenerator:
    """Generate Word documents from pillar analysis JSON files"""
    
    def __init__(self):
        self.docx_available = DOCX_AVAILABLE
        
    def create_document(self, data: Dict) -> Optional[str]:
        """Create a Word document from pillar analysis data"""
        if not self.docx_available:
            return None
        
        try:
            # Extract data with error handling
            metadata = data.get("metadata", {})
            analysis = data.get("analysis", {})
            
            if not metadata:
                return None
            
            # Create document
            doc = Document()
            
            # Set up styles
            self._setup_styles(doc)
            
            # Add content
            if analysis:
                self._add_detailed_analysis(doc, analysis)
                self._add_author_info(doc, metadata)
            else:
                doc.add_paragraph("No analysis data available.")
            
            # Save document
            word_docs_dir = "word_documents"
            if not os.path.exists(word_docs_dir):
                os.makedirs(word_docs_dir)
                
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            pillar = metadata.get('pillar', 'Unknown')
            product = metadata.get('product', 'Unknown')
            pillar_name = pillar.lower().replace(" ", "_")
            product_name = product.lower().replace(" ", "_").replace("temenos_", "")
            filename = f"{pillar_name}_analysis_{product_name}_{timestamp}.docx"
            filepath = os.path.join(word_docs_dir, filename)
            
            doc.save(filepath)
            return filepath
            
        except Exception as e:
            print(f"Error creating Word document: {e}")
            return None
    
    def _setup_styles(self, doc: Document):
        """Set up document styles"""
        try:
            # Title style
            title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Calibri'
            title_style.font.size = Pt(16)
            title_style.font.bold = True
            
            # Heading style
            heading_style = doc.styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.name = 'Calibri'
            heading_style.font.size = Pt(14)
            heading_style.font.bold = True
            
        except Exception:
            pass  # Use default styles if custom styles fail
    
    def _add_detailed_analysis(self, doc: Document, analysis: Dict):
        """Add detailed questions and answers section with integrated business benefits"""
        doc.add_heading('Base Questions & Answers', level=1)
        
        # Check if conversation_flow exists and is not empty
        if 'conversation_flow' not in analysis or not analysis['conversation_flow']:
            doc.add_paragraph("No conversation flow data available.")
            return
        
        # Add questions in compact format
        for i, q_data in enumerate(analysis['conversation_flow'], 1):
            question = q_data.get('question', 'No question available')
            answer = q_data.get('answer', 'No answer available')
            
            # Add question as bold
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"Q{i}: {question}")
            q_run.bold = True
            
            # Add answer with integrated business benefit
            answer_para = doc.add_paragraph(answer)
            answer_para.style = 'Normal'
            
            # Add small spacing between questions
            doc.add_paragraph()
    
    def _add_author_info(self, doc: Document, metadata: Dict):
        """Add author information at the end"""
        # Add page break before author info
        doc.add_page_break()
        
        # Add author section
        doc.add_heading('Document Information', level=1)
        
        # Get API key from shared_config
        try:
            from shared_config import API_CONFIG
            api_key = API_CONFIG.get('jwt_token', 'N/A')
            # Show only first 10 and last 10 characters for security
            if len(api_key) > 20:
                masked_key = f"{api_key[:10]}...{api_key[-10:]}"
            else:
                masked_key = "N/A"
        except:
            masked_key = "N/A"
        
        # Add author info
        doc.add_paragraph(f"Generated by: Temenos RAG AI System")
        doc.add_paragraph(f"API Key: {masked_key}")
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Product: {metadata.get('product', 'Unknown')}")
        doc.add_paragraph(f"Pillar: {metadata.get('pillar', 'Unknown')}")
        doc.add_paragraph(f"Region: {metadata.get('region', 'Unknown')}")
    
    def convert_json_to_word(self, json_filepath: str) -> Optional[str]:
        """Convert JSON analysis file to Word document"""
        try:
            with open(json_filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            return self.create_document(data)
            
        except Exception as e:
            print(f"Error converting JSON to Word: {e}")
            return None
